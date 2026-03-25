"""
modules/docx_builder.py
========================
Builds a formatted .docx file from structured page data using python-docx.

Features
--------
* Three visual themes: academic, minimal, professional
* Auto table-of-contents with page titles
* Per-page sections with clear separators
* Key terms in bold, formulas in monospace
* Confidence score indicator in footer
* Optional original image embedding
* Export to PDF via LibreOffice headless (if installed)

pip install python-docx
"""

from __future__ import annotations

import os
import subprocess
import shutil
import traceback
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx

from config import Config


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _safe_text(text) -> str:
    """
    Strip characters that python-docx / Word XML cannot handle.
    Removes control chars, emojis, and non-BMP Unicode (> U+FFFF)
    which cause UnicodeEncodeError on Windows.
    """
    if not text:
        return ""
    cleaned = []
    for ch in str(text):
        cp = ord(ch)
        # Allow tab(9), LF(10), CR(13), printable ASCII, and Latin/CJK up to U+FFFF
        if cp in (9, 10, 13) or (32 <= cp <= 126) or (160 <= cp <= 0xFFFF):
            cleaned.append(ch)
        else:
            cleaned.append("?")   # safe replacement for emoji / non-BMP
    return "".join(cleaned)


def _conf_label(confidence: float) -> str:
    """Plain-text confidence label — no emojis."""
    if confidence >= 80:
        return "[HIGH]"
    elif confidence >= 50:
        return "[MED]"
    return "[LOW]"


def _add_horizontal_rule(doc: Document):
    """Insert a thin horizontal line via paragraph bottom-border."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "AAAAAA")
    pBdr.append(bot)
    pPr.append(pBdr)


# ─── Main Class ───────────────────────────────────────────────────────────────

class DOCXBuilder:

    def __init__(self):
        self._ready = True

    # ── Public API ────────────────────────────────────────────────────────────

    def build(
        self,
        pages      : list,
        output_dir : str,
        session_id : str,
        settings   : dict,
    ) -> str:
        """
        Build a DOCX from processed page data and return its file path.
        """
        theme_name  = settings.get("theme", Config.DEFAULT_THEME)
        theme       = Config.DOCX_THEMES.get(theme_name, Config.DOCX_THEMES["academic"])
        doc_title   = _safe_text(settings.get("doc_title",   "AI Digitized Notes"))
        doc_subject = _safe_text(settings.get("doc_subject", ""))
        doc_date    = _safe_text(settings.get("doc_date",    ""))
        include_img = settings.get("include_images", False)
        font_name   = settings.get("font",      theme["body_font"])
        font_size   = int(settings.get("font_size", theme["body_size"]))

        doc = Document()

        # Page margins
        sec = doc.sections[0]
        sec.page_height   = Cm(29.7)
        sec.page_width    = Cm(21.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)

        self._create_styles(doc, theme, font_name, font_size)
        self._add_cover(doc, doc_title, doc_subject, doc_date, pages)

        if len(pages) > 1:
            self._add_toc(doc, pages)

        for idx, page in enumerate(pages):
            try:
                self._add_page_section(
                    doc, page, idx, theme, font_name, font_size, include_img
                )
            except Exception as e:
                print(f"[DOCX] Error rendering page {idx}: {e}")
                traceback.print_exc()
                doc.add_paragraph(
                    f"Page {idx + 1} — rendering error: {_safe_text(str(e))}",
                    style="NoteBody"
                )
                doc.add_page_break()

        os.makedirs(output_dir, exist_ok=True)
        out_path = os.path.join(output_dir, f"notes_{session_id[:8]}.docx")

        try:
            doc.save(out_path)
            print(f"[DOCX] Saved: {out_path}")
            return out_path
        except Exception as e:
            print(f"[DOCX] Primary save failed: {e}. Trying fallback path...")
            traceback.print_exc()
            fallback = os.path.join(output_dir, f"notes_{session_id[:8]}_fallback.docx")
            doc.save(fallback)
            print(f"[DOCX] Fallback saved: {fallback}")
            return fallback

    def export_pdf(self, docx_path: str) -> Optional[str]:
        """Convert DOCX to PDF using LibreOffice headless."""
        lo_cmd = shutil.which("libreoffice") or shutil.which("soffice")
        if not lo_cmd:
            print("[DOCX] LibreOffice not found — cannot export PDF.")
            print("       Windows: https://www.libreoffice.org/download/download-libreoffice/")
            return None

        out_dir = str(Path(docx_path).parent)
        try:
            result = subprocess.run(
                [lo_cmd, "--headless", "--convert-to", "pdf",
                 "--outdir", out_dir, docx_path],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode != 0:
                print(f"[DOCX] LibreOffice error: {result.stderr}")
                return None
            pdf_path = str(Path(docx_path).with_suffix(".pdf"))
            if os.path.isfile(pdf_path):
                print(f"[DOCX] PDF exported: {pdf_path}")
                return pdf_path
        except subprocess.TimeoutExpired:
            print("[DOCX] LibreOffice timed out")
        except Exception as e:
            print(f"[DOCX] PDF export error: {e}")
        return None

    def is_ready(self) -> bool:
        return self._ready

    # ── Styles ────────────────────────────────────────────────────────────────

    def _create_styles(self, doc, theme, font_name, font_size):
        styles = doc.styles

        def get_or_add(name):
            try:
                return styles[name]
            except KeyError:
                return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

        # NoteTitle
        s = get_or_add("NoteTitle")
        s.font.name = theme["title_font"]
        s.font.size = Pt(theme["title_size"])
        s.font.bold = True
        s.font.color.rgb = _hex_to_rgb(theme["title_color"])
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after  = Pt(6)

        # NoteHeading1
        s = get_or_add("NoteHeading1")
        s.font.name = theme["title_font"]
        s.font.size = Pt(theme["heading_size"])
        s.font.bold = True
        s.font.color.rgb = _hex_to_rgb(theme["heading_color"])
        s.paragraph_format.space_before = Pt(8)
        s.paragraph_format.space_after  = Pt(4)

        # NoteHeading2
        s = get_or_add("NoteHeading2")
        s.font.name   = theme["title_font"]
        s.font.size   = Pt(theme["heading_size"] - 1)
        s.font.bold   = True
        s.font.italic = True
        s.font.color.rgb = _hex_to_rgb(theme["heading_color"])
        s.paragraph_format.space_before = Pt(6)
        s.paragraph_format.space_after  = Pt(3)

        # NoteBody
        s = get_or_add("NoteBody")
        s.font.name = font_name
        s.font.size = Pt(font_size)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        s.paragraph_format.line_spacing      = 1.15

        # NoteBullet
        s = get_or_add("NoteBullet")
        s.font.name = font_name
        s.font.size = Pt(font_size)
        s.paragraph_format.left_indent = Cm(0.7)
        s.paragraph_format.space_after = Pt(2)

        # NoteFormula
        s = get_or_add("NoteFormula")
        s.font.name = "Courier New"
        s.font.size = Pt(font_size + 1)
        s.font.bold = True
        s.font.color.rgb = _hex_to_rgb(theme["accent_color"])
        s.paragraph_format.left_indent  = Cm(0.5)
        s.paragraph_format.space_before = Pt(3)
        s.paragraph_format.space_after  = Pt(3)

        # NoteSummary
        s = get_or_add("NoteSummary")
        s.font.name   = font_name
        s.font.size   = Pt(font_size - 1)
        s.font.italic = True
        s.paragraph_format.left_indent = Cm(0.3)
        s.paragraph_format.space_after = Pt(6)

        # NoteCaption
        s = get_or_add("NoteCaption")
        s.font.name      = font_name
        s.font.size      = Pt(8)
        s.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        s.paragraph_format.space_after = Pt(8)

    # ── Cover Page ────────────────────────────────────────────────────────────

    def _add_cover(self, doc, title, subject, date, pages):
        p = doc.add_paragraph(_safe_text(title) or "AI Digitized Notes",
                               style="NoteTitle")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if subject:
            p = doc.add_paragraph(_safe_text(subject), style="NoteHeading2")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if date:
            p = doc.add_paragraph(_safe_text(date), style="NoteCaption")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subjects = list({pg.get("subject", "General") for pg in pages if pg.get("subject")})
        if subjects:
            p = doc.add_paragraph(
                "Subjects: " + ", ".join(_safe_text(s) for s in subjects),
                style="NoteSummary"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph(f"Total Pages: {len(pages)}", style="NoteCaption")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph("Generated by AI Notes Digitizer", style="NoteCaption")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _add_horizontal_rule(doc)
        doc.add_page_break()

    # ── Table of Contents ─────────────────────────────────────────────────────

    def _add_toc(self, doc, pages):
        doc.add_paragraph("Table of Contents", style="NoteTitle")
        _add_horizontal_rule(doc)

        for idx, page in enumerate(pages):
            fmt   = page.get("formatted", {})
            title = _safe_text(fmt.get("title") or page.get("page_label", f"Page {idx + 1}"))
            subj  = _safe_text(page.get("subject", ""))
            label = f"{idx + 1}.  {title}"
            if subj:
                label += f"  [{subj}]"
            p = doc.add_paragraph(label, style="NoteBody")
            p.paragraph_format.space_after = Pt(2)

        doc.add_page_break()

    # ── Per-page Section ──────────────────────────────────────────────────────

    def _add_page_section(self, doc, page, idx, theme, font_name, font_size, include_img):
        fmt        = page.get("formatted", {})
        confidence = float(page.get("confidence", 0))
        subject    = _safe_text(page.get("subject",    "General"))
        page_label = _safe_text(page.get("page_label", f"Page {idx + 1}"))

        # Section title
        title = _safe_text(fmt.get("title") or page_label)
        doc.add_paragraph(f"{idx + 1}. {title}", style="NoteHeading1")

        # Metadata
        meta = doc.add_paragraph(style="NoteCaption")
        meta.add_run(
            f"Subject: {subject}  |  "
            f"OCR Confidence: {confidence:.0f}% {_conf_label(confidence)}  |  "
            f"Source: {page_label}"
        )

        # Summary
        summary = _safe_text(fmt.get("summary", "")).strip()
        if summary:
            doc.add_paragraph("Summary", style="NoteHeading2")
            doc.add_paragraph(summary, style="NoteSummary")

        # Key terms
        key_terms = [_safe_text(t) for t in fmt.get("key_terms", []) if t]
        if key_terms:
            kt = doc.add_paragraph(style="NoteBody")
            r  = kt.add_run("Key Terms: ")
            r.bold = True
            for i, term in enumerate(key_terms[:10]):
                r = kt.add_run(term)
                r.bold = True
                r.font.color.rgb = _hex_to_rgb(theme["accent_color"])
                if i < len(key_terms) - 1:
                    kt.add_run(", ")

        _add_horizontal_rule(doc)

        # Original image (optional)
        if include_img:
            img_path = page.get("enhanced_image")
            if img_path and os.path.isfile(str(img_path)):
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                    cap = doc.add_paragraph(
                        f"[Original scanned image - {page_label}]",
                        style="NoteCaption"
                    )
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_horizontal_rule(doc)
                except Exception as e:
                    print(f"[DOCX] Could not insert image for page {idx}: {e}")

        # Sections
        sections = fmt.get("sections", [])
        if not sections:
            raw = _safe_text(page.get("raw_text", ""))
            if raw:
                doc.add_paragraph(raw, style="NoteBody")
        else:
            for sec in sections:
                heading  = _safe_text(sec.get("heading",  "")).strip()
                content  = _safe_text(sec.get("content",  "")).strip()
                bullets  = [_safe_text(b) for b in sec.get("bullets",  []) if b]
                formulas = [_safe_text(f) for f in sec.get("formulas", []) if f]
                notes    = _safe_text(sec.get("notes", "")).strip()

                if heading:
                    doc.add_paragraph(heading, style="NoteHeading2")
                if content:
                    doc.add_paragraph(content, style="NoteBody")
                for bullet in bullets:
                    p = doc.add_paragraph(style="NoteBullet")
                    p.add_run("* ").bold = True
                    p.add_run(bullet)
                for formula in formulas:
                    p   = doc.add_paragraph(style="NoteFormula")
                    run = p.add_run(f"  {formula}")
                    run.font.name      = "Courier New"
                    run.font.size      = Pt(font_size + 1)
                    run.font.bold      = True
                    run.font.color.rgb = _hex_to_rgb(theme["accent_color"])
                if notes:
                    p = doc.add_paragraph(style="NoteCaption")
                    p.add_run(f"Note: {notes}").italic = True

        # Formulas block
        all_formulas = page.get("formulas", [])
        if all_formulas:
            doc.add_paragraph("Formulas Detected", style="NoteHeading2")
            for f in all_formulas:
                uni   = _safe_text(f.get("unicode", ""))
                ftype = _safe_text(f.get("type",    "inline"))
                doc.add_paragraph(f"  {uni}  [{ftype}]", style="NoteFormula")

        doc.add_page_break()